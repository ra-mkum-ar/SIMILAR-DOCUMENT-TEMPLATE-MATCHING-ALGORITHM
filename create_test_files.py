#!/usr/bin/env python3
"""
Create sample PDF and DOCX files for testing the document matching application
"""

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
import os

def create_sample_pdf(filename, title, content):
    """Create a sample PDF file"""
    c = canvas.Canvas(filename, pagesize=letter)
    width, height = letter
    
    # Title
    c.setFont("Helvetica-Bold", 16)
    c.drawString(50, height - 50, title)
    
    # Content
    c.setFont("Helvetica", 12)
    y_position = height - 100
    
    for line in content.split('\n'):
        if y_position < 50:  # Start new page if needed
            c.showPage()
            y_position = height - 50
        c.drawString(50, y_position, line)
        y_position -= 20
    
    c.save()

def create_sample_docx(filename, title, content):
    """Create a sample DOCX file"""
    doc = Document()
    
    # Add title
    title_para = doc.add_heading(title, 0)
    
    # Add content
    for paragraph in content.split('\n\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    
    doc.save(filename)

# Create test files directory
os.makedirs('/app/test_files', exist_ok=True)

# Template 1: Business Contract
contract_content = """
BUSINESS SERVICE AGREEMENT

This Service Agreement ("Agreement") is entered into on [DATE] between [COMPANY A] and [COMPANY B].

1. SCOPE OF SERVICES
The Provider agrees to deliver the following services:
- Software development and maintenance
- Technical support and consultation
- Project management services

2. PAYMENT TERMS
Payment shall be made within 30 days of invoice receipt.
Total project cost: $50,000
Payment schedule: 50% upfront, 50% upon completion

3. DELIVERABLES
- Complete software application
- Documentation and user manuals
- Training materials

4. TIMELINE
Project duration: 6 months
Milestone reviews: Monthly
Final delivery: [END DATE]

5. CONFIDENTIALITY
Both parties agree to maintain confidentiality of proprietary information.

6. TERMINATION
Either party may terminate with 30 days written notice.
"""

# Template 2: Technical Manual
manual_content = """
SOFTWARE INSTALLATION GUIDE

SYSTEM REQUIREMENTS
- Operating System: Windows 10 or later
- RAM: 8GB minimum, 16GB recommended
- Storage: 500MB available space
- Network: Internet connection required

INSTALLATION STEPS

Step 1: Download Software
Visit our website and download the latest version.
Verify the file integrity using the provided checksum.

Step 2: Run Installer
Double-click the installer file.
Follow the on-screen instructions.
Accept the license agreement.

Step 3: Configuration
Enter your license key when prompted.
Configure network settings.
Set up user preferences.

Step 4: Verification
Launch the application.
Run the built-in diagnostic tool.
Verify all features are working correctly.

TROUBLESHOOTING

Common Issues:
- Installation fails: Check system requirements
- License error: Verify key format
- Network issues: Check firewall settings

For additional support, contact our technical team.
"""

# Template 3: Research Report
research_content = """
MARKET ANALYSIS REPORT

EXECUTIVE SUMMARY
This report analyzes current market trends in the technology sector.
Key findings indicate significant growth in AI and cloud computing.

METHODOLOGY
Data collection period: January 2024 - June 2024
Sample size: 500 companies
Research methods: Surveys, interviews, financial analysis

KEY FINDINGS

Market Growth
- AI sector: 45% year-over-year growth
- Cloud computing: 32% growth
- Cybersecurity: 28% growth

Investment Trends
Total investment in tech startups: $12.5 billion
Top funded categories:
1. Artificial Intelligence
2. Fintech
3. Healthcare Technology

RECOMMENDATIONS
1. Increase investment in AI research
2. Develop cloud-native solutions
3. Strengthen cybersecurity measures

CONCLUSION
The technology sector shows strong growth potential.
Companies should focus on emerging technologies.
Strategic partnerships will be crucial for success.
"""

# Query documents for testing
query1_content = """
SERVICE AGREEMENT TEMPLATE

This Agreement is between ABC Corp and XYZ Ltd for software development services.

SERVICES PROVIDED:
- Custom application development
- System integration
- Ongoing maintenance and support

FINANCIAL TERMS:
Total cost: $45,000
Payment: 40% upfront, 60% on delivery
Invoice terms: Net 30 days

PROJECT SCOPE:
- Web application development
- Mobile app creation  
- User training and documentation

TIMELINE:
Duration: 5 months
Reviews: Bi-weekly
Completion: December 2024

LEGAL TERMS:
Confidentiality clause included
30-day termination notice required
"""

query2_content = """
INSTALLATION MANUAL

SYSTEM PREREQUISITES
- Windows 11 or macOS 12+
- 16GB RAM recommended
- 1GB free disk space
- Stable internet connection

SETUP PROCESS

Phase 1: Preparation
Download the software package
Verify system compatibility
Backup existing data

Phase 2: Installation
Execute the setup file
Follow installation wizard
Enter product activation key

Phase 3: Initial Setup
Configure user accounts
Set network parameters
Customize interface preferences

Phase 4: Testing
Run system diagnostics
Test core functionality
Validate installation success

SUPPORT INFORMATION
For technical assistance, visit our help center
Email: support@company.com
Phone: 1-800-SUPPORT
"""

# Create template files
print("Creating template files...")
create_sample_pdf('/app/test_files/template_contract.pdf', 'Business Service Agreement', contract_content)
create_sample_docx('/app/test_files/template_manual.docx', 'Software Installation Guide', manual_content)
create_sample_pdf('/app/test_files/template_research.pdf', 'Market Analysis Report', research_content)

# Create query files for testing
print("Creating query files...")
create_sample_docx('/app/test_files/query_agreement.docx', 'Service Agreement Template', query1_content)
create_sample_pdf('/app/test_files/query_installation.pdf', 'Installation Manual', query2_content)

print("Test files created successfully!")
print("Template files:")
print("- /app/test_files/template_contract.pdf")
print("- /app/test_files/template_manual.docx") 
print("- /app/test_files/template_research.pdf")
print("Query files:")
print("- /app/test_files/query_agreement.docx")
print("- /app/test_files/query_installation.pdf")